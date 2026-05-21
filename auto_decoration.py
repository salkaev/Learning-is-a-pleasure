from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

table = None

# Открываем исходный документ
doc = Document("1000.docx")

# Создаём новый документ
new_doc = Document()

# Настройка полей в НОВОМ документе
section = new_doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(3)
section.right_margin = Cm(1.5)
FLAG = False
cnt = 0
for paragraph in doc.paragraphs:
    # Если это ГЛАВА
    if "Глава" in paragraph.text:
        cnt += 1
        # Добавляем разрыв страницы ПЕРЕД главой
        if cnt > 2 :
            new_doc.add_page_break()

        # Добавляем главу в новый документ
        new_paragraph = new_doc.add_paragraph(paragraph.text)
        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_paragraph.runs[0].font.size = Pt(16)
        new_paragraph.runs[0].font.bold = True
        new_paragraph.runs[0].font.name = 'Times New Roman'

        # Настройка отступов для главы
        fmt = new_paragraph.paragraph_format
        fmt.line_spacing = 1.5
        fmt.first_line_indent = Cm(0)
        
    elif "```" in paragraph.text or FLAG == True:
        # Если это открывающие ```, создаём новую таблицу
        if "```" in paragraph.text and FLAG == False:
            table = new_doc.add_table(rows=1, cols=1)
            # ← ИСПРАВЛЕНИЕ: настройка шрифта для ячейки таблицы
            for paragraph_in_cell in table.cell(0,0).paragraphs:
                for run in paragraph_in_cell.runs:
                    run.font.size = Pt(14)
                    run.font.name = 'Times New Roman'
            FLAG = True
            continue
            
        # Если это закрывающие ```, завершаем листинг
        elif "```" in paragraph.text and FLAG == True:
            FLAG = False
            continue
            
        # Добавляем код в таблицу
        if FLAG == True:
            table.cell(0,0).text += paragraph.text + "\n"
            # ← ИСПРАВЛЕНИЕ: применяем шрифт после добавления текста
            for paragraph_in_cell in table.cell(0,0).paragraphs:
                for run in paragraph_in_cell.runs:
                    run.font.size = Pt(14)
                    run.font.name = 'Times New Roman'

    else:
        # Пропускаем пустые абзацы только если не в режиме листинга
        if not paragraph.text.strip():
            continue
            
        # Добавляем обычный текст в новый документ
        new_paragraph = new_doc.add_paragraph(paragraph.text)
        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        new_paragraph.runs[0].font.size = Pt(14)
        new_paragraph.runs[0].font.name = 'Times New Roman'

        # Настройка отступов для обычного текста
        fmt = new_paragraph.paragraph_format
        fmt.line_spacing = 1.5
        fmt.first_line_indent = Cm(1.25)

# Сохраняем НОВЫЙ документ (исходный не трогаем)
new_doc.save("1000_formatted.docx")
print(f"Готово! Найдено глав: {cnt}")
print("Сохранено как: 1000_formatted.docx")